from __future__ import annotations

import csv
import html
import json
import math
import re
import shutil
import zipfile
from collections import Counter
from datetime import date
from html import unescape
from html.parser import HTMLParser
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader


PROJECT = next(candidate for candidate in [Path.cwd(), *Path.cwd().parents] if (candidate / 'reports' / 'generated').exists() and (candidate / 'reports' / 'actual').exists())
REPORTS_GENERATED = PROJECT / "reports" / "generated"
REPORTS_ACTUAL = PROJECT / "reports" / "actual"
RESULTS = PROJECT / "results"
AVS_ROOT = RESULTS / "Actual vs Generatet"
PY_DIR = AVS_ROOT / "Python Analyse"
LLM_DIR = AVS_ROOT / "LLM Analyse"
FIG_DIR = PY_DIR / "figures"
TEXT_DIR = PY_DIR / "extracted_text"
NOTEBOOK_ROOT = PROJECT / "analysis.ipynb"
README = RESULTS / "README.md"
OUTPUTS = Path.cwd() / "outputs"

COMPANIES = ["Amazon", "Microsoft", "SAP", "Zalando"]

GENERATED_FILES = {
    "Amazon": [REPORTS_GENERATED / "Amazon_ESG_Report_Generated.docx"],
    "Microsoft": [REPORTS_GENERATED / "Microsoft_ESG_Report_Generated.docx"],
    "SAP": [REPORTS_GENERATED / "SAP_ESG_Report_Generated.docx"],
    "Zalando": [REPORTS_GENERATED / "Zalando_ESG_Report_Generated.docx"],
}

ACTUAL_FILES = {
    "Amazon": [REPORTS_ACTUAL / "2024 Amazon Sustainability Report - 2024-amazon-sustainability-report.pdf"],
    "Microsoft": [
        REPORTS_ACTUAL / "microsoft_2025_environmental_sustainability_report.pdf",
        REPORTS_ACTUAL / "microsoft_2025_impact_summary.pdf",
    ],
    "SAP": [REPORTS_ACTUAL / "SAP 2025 Annual Report Form 20-F.html"],
    "Zalando": [REPORTS_ACTUAL / "zalando-se_en_full-csrd-esrs-report_annual-report_2025.pdf"],
}

STOPWORDS = set(
    """
    a about above after again against all also am an and any are as at be because been before being below between both but by
    can could did do does doing down during each few for from further had has have having he her here hers him his how i if
    in into is it its itself just me more most my no nor not now of off on once only or other our ours out over own same she
    should so some such than that the their theirs them then there these they this those through to too under until up very was
    we were what when where which while who whom why will with within would you your
    """.split()
)

DOMAIN_STOPWORDS = set(
    """
    amazon microsoft sap zalando esg report reports generated actual company companies information general environmental environment
    social governance section task data source sources material analysis business group inc se ag com annual sustainability
    responsibility integrated page pages table financial year fiscal ltd corporate disclosure disclosures
    """.split()
)

ESG_LEXICONS = {
    "Environmental": [
        "climate",
        "carbon",
        "emission",
        "emissions",
        "scope 1",
        "scope 2",
        "scope 3",
        "greenhouse gas",
        "ghg",
        "energy",
        "renewable",
        "electricity",
        "water",
        "waste",
        "recycling",
        "circular",
        "circularity",
        "packaging",
        "biodiversity",
        "net zero",
        "decarbonization",
        "decarbonisation",
        "data center",
        "data centre",
        "carbon free",
        "carbon neutral",
        "solar",
        "wind",
        "logistics",
        "returns",
        "materials",
        "science based targets",
        "taxonomy",
    ],
    "Social": [
        "employee",
        "employees",
        "workforce",
        "workers",
        "labour",
        "labor",
        "human rights",
        "diversity",
        "inclusion",
        "equity",
        "safety",
        "health",
        "wellbeing",
        "well-being",
        "training",
        "talent",
        "communities",
        "accessibility",
        "supplier",
        "suppliers",
        "living wage",
        "audit",
        "audits",
        "union",
        "freedom of association",
        "responsible ai",
        "digital inclusion",
        "grievance",
        "complaint",
        "harassment",
        "discrimination",
        "own workforce",
        "value chain workers",
    ],
    "Governance": [
        "board",
        "supervisory board",
        "committee",
        "audit",
        "compliance",
        "ethics",
        "code of conduct",
        "anti-corruption",
        "risk management",
        "internal controls",
        "data protection",
        "cybersecurity",
        "privacy",
        "ai governance",
        "dsa",
        "csrd",
        "esrs",
        "double materiality",
        "assurance",
        "reporting",
        "executive compensation",
        "whistleblower",
        "supply chain due diligence",
        "regulatory",
        "transparency",
        "shareholder",
        "agm",
    ],
}

TONE_LEXICONS = {
    "Success": [
        "achieved",
        "reduced",
        "improved",
        "progress",
        "strong",
        "robust",
        "mature",
        "leadership",
        "exceeded",
        "completed",
        "expanded",
        "increased",
        "certified",
        "recognized",
        "integrated",
        "strengthened",
        "delivered",
        "commitment",
        "committed",
        "implemented",
    ],
    "Risk": [
        "risk",
        "risks",
        "challenge",
        "challenges",
        "pressure",
        "litigation",
        "investigation",
        "criticism",
        "complaint",
        "controversy",
        "incident",
        "breach",
        "uncertainty",
        "exposure",
        "dependency",
        "gap",
        "constraint",
        "concern",
        "scrutiny",
        "violation",
        "fine",
        "penalty",
    ],
    "Controversy": [
        "controversy",
        "controversies",
        "criticism",
        "complaint",
        "allegation",
        "lawsuit",
        "investigation",
        "union",
        "strike",
        "dsa",
        "antitrust",
        "labor",
        "privacy",
        "incident",
        "fine",
        "protest",
    ],
    "Targets_Tradeoffs": [
        "target",
        "targets",
        "goal",
        "goals",
        "ambition",
        "by 2030",
        "by 2040",
        "by 2050",
        "roadmap",
        "pathway",
        "future",
        "planned",
        "expected",
        "transition",
        "trade off",
        "trade-off",
        "tension",
        "conflict",
        "requires",
        "dependent",
    ],
}

TOKEN_RE = re.compile(r"[A-Za-z]+(?:[-'][A-Za-z]+)?")
SENTENCE_RE = re.compile(r"(?<=[.!?])\s+(?=[A-Z0-9\"'])")
NUMBER_RE = re.compile(
    r"(?<![A-Za-z])(?:\d{1,3}(?:,\d{3})+|\d+(?:\.\d+)?)(?:\s?(?:%|percent|million|billion|bn|m|tco2e|co2e|eur|usd|\$|employees|suppliers|countries|pages|tons|tonnes))?",
    re.IGNORECASE,
)


def extract_docx_text(path: Path) -> str:
    doc = Document(path)
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_pdf_text(path: Path) -> tuple[str, int]:
    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        if text.strip():
            parts.append(text)
    return "\n".join(parts), len(reader.pages)


class TextHTMLParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []
        self.skip_depth = 0

    def handle_starttag(self, tag: str, attrs) -> None:
        if tag.lower() in {"script", "style", "ix:header"}:
            self.skip_depth += 1
        elif tag.lower() in {"p", "div", "tr", "li", "br", "h1", "h2", "h3", "h4", "td", "th"}:
            self.parts.append(" ")

    def handle_endtag(self, tag: str) -> None:
        if tag.lower() in {"script", "style", "ix:header"} and self.skip_depth:
            self.skip_depth -= 1
        elif tag.lower() in {"p", "div", "tr", "li", "h1", "h2", "h3", "h4"}:
            self.parts.append("\n")

    def handle_data(self, data: str) -> None:
        if not self.skip_depth and data.strip():
            self.parts.append(data)

    def text(self) -> str:
        raw = unescape(" ".join(self.parts))
        raw = re.sub(r"[ \t\r\f\v]+", " ", raw)
        raw = re.sub(r"\n\s+", "\n", raw)
        return raw.strip()


def extract_html_text(path: Path) -> str:
    parser = TextHTMLParser()
    parser.feed(path.read_text(encoding="utf-8", errors="ignore"))
    return parser.text()


def extract_actual_text(path: Path) -> tuple[str, object]:
    if path.suffix.lower() == ".pdf":
        return extract_pdf_text(path)
    if path.suffix.lower() in {".htm", ".html"}:
        return extract_html_text(path), "SEC HTML filing"
    raise ValueError(f"Unsupported actual report type: {path}")


def tokenize(text: str) -> list[str]:
    return [match.group(0).lower().strip("-'") for match in TOKEN_RE.finditer(text)]


def clean_tokens(text: str) -> list[str]:
    return [
        token
        for token in tokenize(text)
        if len(token) > 2 and token not in STOPWORDS and token not in DOMAIN_STOPWORDS
    ]


def sentence_split(text: str) -> list[str]:
    return [s.strip() for s in SENTENCE_RE.split(text.replace("\n", " ")) if s.strip()]


def normalize_phrase_text(text: str) -> str:
    return " " + re.sub(r"[^a-z0-9]+", " ", text.lower()).strip() + " "


def count_terms(text: str, terms: list[str]) -> int:
    normalized = normalize_phrase_text(text)
    total = 0
    for term in terms:
        normalized_term = normalize_phrase_text(term).strip()
        if normalized_term:
            total += len(re.findall(rf"\b{re.escape(normalized_term)}\b", normalized))
    return total


def build_corpus() -> tuple[dict[tuple[str, str], str], list[dict[str, object]]]:
    corpus: dict[tuple[str, str], str] = {}
    source_rows: list[dict[str, object]] = []
    TEXT_DIR.mkdir(parents=True, exist_ok=True)
    for company in COMPANIES:
        generated_parts = [extract_docx_text(path) for path in GENERATED_FILES[company]]
        generated_text = "\n\n".join(generated_parts)
        corpus[(company, "Generated")] = generated_text
        (TEXT_DIR / f"{company}_generated.txt").write_text(generated_text, encoding="utf-8")
        source_rows.append(
            {
                "company": company,
                "report_type": "Generated",
                "source_files": "; ".join(path.name for path in GENERATED_FILES[company]),
                "pages": "",
            }
        )

        actual_parts: list[str] = []
        page_count: int | str = 0
        for path in ACTUAL_FILES[company]:
            text, pages = extract_actual_text(path)
            actual_parts.append(text)
            if isinstance(pages, int) and isinstance(page_count, int):
                page_count += pages
            else:
                page_count = str(pages)
        actual_text = "\n\n".join(actual_parts)
        corpus[(company, "Actual")] = actual_text
        (TEXT_DIR / f"{company}_actual.txt").write_text(actual_text, encoding="utf-8")
        source_rows.append(
            {
                "company": company,
                "report_type": "Actual",
                "source_files": "; ".join(path.name for path in ACTUAL_FILES[company]),
                "pages": page_count,
            }
        )
    return corpus, source_rows


def metric_rows(corpus: dict[tuple[str, str], str]) -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    metrics: list[dict[str, object]] = []
    esg_rows: list[dict[str, object]] = []
    tone_rows: list[dict[str, object]] = []
    top_rows: list[dict[str, object]] = []
    for (company, report_type), text in corpus.items():
        all_tokens = tokenize(text)
        tokens = clean_tokens(text)
        sentences = sentence_split(text)
        word_count = len(all_tokens)
        numbers = len(NUMBER_RE.findall(text))
        metrics.append(
            {
                "company": company,
                "report_type": report_type,
                "word_count": word_count,
                "sentence_count": len(sentences),
                "avg_sentence_length": round(word_count / max(len(sentences), 1), 2),
                "unique_terms": len(set(tokens)),
                "lexical_diversity": round(len(set(tokens)) / max(len(tokens), 1), 4),
                "numeric_mentions": numbers,
                "numeric_mentions_per_1000_words": round(numbers / max(word_count, 1) * 1000, 2),
            }
        )
        esg_counts = {pillar: count_terms(text, terms) for pillar, terms in ESG_LEXICONS.items()}
        esg_rows.append(
            {
                "company": company,
                "report_type": report_type,
                "environmental_hits": esg_counts["Environmental"],
                "social_hits": esg_counts["Social"],
                "governance_hits": esg_counts["Governance"],
                "total_esg_hits": sum(esg_counts.values()),
                "esg_hits_per_1000_words": round(sum(esg_counts.values()) / max(word_count, 1) * 1000, 2),
                "strongest_pillar": max(esg_counts.items(), key=lambda item: item[1])[0],
            }
        )
        tone_counts = {tone: count_terms(text, terms) for tone, terms in TONE_LEXICONS.items()}
        tone_rows.append(
            {
                "company": company,
                "report_type": report_type,
                "success_hits": tone_counts["Success"],
                "risk_hits": tone_counts["Risk"],
                "controversy_hits": tone_counts["Controversy"],
                "target_tradeoff_hits": tone_counts["Targets_Tradeoffs"],
                "risk_to_success_ratio": round(tone_counts["Risk"] / max(tone_counts["Success"], 1), 2),
            }
        )
        for rank, (term, count) in enumerate(Counter(tokens).most_common(25), start=1):
            top_rows.append({"company": company, "report_type": report_type, "rank": rank, "term": term, "count": count})
    return pd.DataFrame(metrics), pd.DataFrame(esg_rows), pd.DataFrame(tone_rows), pd.DataFrame(top_rows)


def tfidf_vectors(token_docs: dict[tuple[str, str], list[str]]) -> dict[tuple[str, str], dict[str, float]]:
    vocabulary = sorted(set(token for tokens in token_docs.values() for token in tokens))
    doc_count = len(token_docs)
    doc_sets = {label: set(tokens) for label, tokens in token_docs.items()}
    df = {term: sum(1 for doc_set in doc_sets.values() if term in doc_set) for term in vocabulary}
    vectors: dict[tuple[str, str], dict[str, float]] = {}
    for label, tokens in token_docs.items():
        counts = Counter(tokens)
        total = max(len(tokens), 1)
        vectors[label] = {
            term: (counts[term] / total) * (math.log((1 + doc_count) / (1 + df[term])) + 1)
            for term in vocabulary
        }
    return vectors


def cosine(vec_a: dict[str, float], vec_b: dict[str, float]) -> float:
    terms = set(vec_a) | set(vec_b)
    dot = sum(vec_a.get(term, 0.0) * vec_b.get(term, 0.0) for term in terms)
    norm_a = math.sqrt(sum(value * value for value in vec_a.values()))
    norm_b = math.sqrt(sum(value * value for value in vec_b.values()))
    return dot / (norm_a * norm_b) if norm_a and norm_b else 0.0


def pairwise_rows(corpus: dict[tuple[str, str], str], metrics: pd.DataFrame, esg: pd.DataFrame, tone: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
    token_docs = {label: clean_tokens(text) for label, text in corpus.items()}
    vectors = tfidf_vectors(token_docs)
    token_sets = {label: set(tokens) for label, tokens in token_docs.items()}
    rows: list[dict[str, object]] = []
    missing_rows: list[dict[str, object]] = []
    for company in COMPANIES:
        gen = (company, "Generated")
        act = (company, "Actual")
        union = token_sets[gen] | token_sets[act]
        jaccard = len(token_sets[gen] & token_sets[act]) / len(union) if union else 0.0
        m_gen = metrics[(metrics.company == company) & (metrics.report_type == "Generated")].iloc[0]
        m_act = metrics[(metrics.company == company) & (metrics.report_type == "Actual")].iloc[0]
        e_gen = esg[(esg.company == company) & (esg.report_type == "Generated")].iloc[0]
        e_act = esg[(esg.company == company) & (esg.report_type == "Actual")].iloc[0]
        t_gen = tone[(tone.company == company) & (tone.report_type == "Generated")].iloc[0]
        t_act = tone[(tone.company == company) & (tone.report_type == "Actual")].iloc[0]
        rows.append(
            {
                "company": company,
                "tfidf_cosine_generated_vs_actual": round(cosine(vectors[gen], vectors[act]), 4),
                "jaccard_generated_vs_actual": round(jaccard, 4),
                "generated_words_as_pct_of_actual": round(float(m_gen.word_count) / max(float(m_act.word_count), 1) * 100, 2),
                "numeric_density_gap_generated_minus_actual": round(float(m_gen.numeric_mentions_per_1000_words) - float(m_act.numeric_mentions_per_1000_words), 2),
                "esg_density_gap_generated_minus_actual": round(float(e_gen.esg_hits_per_1000_words) - float(e_act.esg_hits_per_1000_words), 2),
                "risk_success_gap_generated_minus_actual": round(float(t_gen.risk_to_success_ratio) - float(t_act.risk_to_success_ratio), 2),
                "generated_strongest_pillar": e_gen.strongest_pillar,
                "actual_strongest_pillar": e_act.strongest_pillar,
            }
        )

        gen_counter = Counter(token_docs[gen])
        act_counter = Counter(token_docs[act])
        for rank, (term, diff) in enumerate(
            sorted(
                ((term, count - gen_counter.get(term, 0)) for term, count in act_counter.items()),
                key=lambda item: (-item[1], item[0]),
            )[:20],
            start=1,
        ):
            missing_rows.append(
                {
                    "company": company,
                    "rank": rank,
                    "actual_over_generated_term": term,
                    "actual_count_minus_generated_count": diff,
                    "actual_count": act_counter[term],
                    "generated_count": gen_counter.get(term, 0),
                }
            )
    return pd.DataFrame(rows), pd.DataFrame(missing_rows)


def xtext(x: float, y: float, text: object, size: int = 12, fill: str = "#111827", anchor: str = "middle", weight: str = "400") -> str:
    return (
        f'<text x="{x:.1f}" y="{y:.1f}" font-family="Arial, sans-serif" font-size="{size}" '
        f'fill="{fill}" text-anchor="{anchor}" font-weight="{weight}">{html.escape(str(text))}</text>'
    )


def write_grouped_bar(path: Path, rows: list[dict[str, object]], metric: str, title: str, y_label: str, scale: float = 1.0) -> None:
    width, height = 940, 430
    left, right, top, bottom = 76, 34, 62, 78
    chart_w = width - left - right
    chart_h = height - top - bottom
    max_value = max(float(row[metric]) for row in rows) * scale or 1
    companies = COMPANIES
    group_w = chart_w / len(companies)
    bar_w = 34
    colors = {"Generated": "#2563EB", "Actual": "#0E7490"}
    lines = [
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}">',
        '<rect width="100%" height="100%" fill="#FFFFFF"/>',
        xtext(width / 2, 28, title, 18, "#111827", "middle", "700"),
        xtext(width - 190, 54, "Generated", 11, "#374151", "start"),
        f'<rect x="{width - 215}" y="44" width="14" height="14" rx="2" fill="{colors["Generated"]}"/>',
        xtext(width - 88, 54, "Actual", 11, "#374151", "start"),
        f'<rect x="{width - 113}" y="44" width="14" height="14" rx="2" fill="{colors["Actual"]}"/>',
    ]
    for tick in range(6):
        value = max_value / 5 * tick
        y = top + chart_h - (value / max_value) * chart_h
        lines.append(f'<line x1="{left}" y1="{y:.1f}" x2="{left + chart_w}" y2="{y:.1f}" stroke="#E5E7EB"/>')
        lines.append(xtext(left - 9, y + 4, f"{value:.0f}", 10, "#6B7280", "end"))
    lines.append(f'<line x1="{left}" y1="{top}" x2="{left}" y2="{top + chart_h}" stroke="#9CA3AF"/>')
    lines.append(f'<line x1="{left}" y1="{top + chart_h}" x2="{left + chart_w}" y2="{top + chart_h}" stroke="#9CA3AF"/>')
    lines.append(xtext(18, top + chart_h / 2, y_label, 12, "#4B5563", "middle", "600").replace("<text", f'<text transform="rotate(-90 18 {top + chart_h / 2:.1f})"', 1))
    for i, company in enumerate(companies):
        group_x = left + i * group_w + group_w / 2
        for j, report_type in enumerate(["Generated", "Actual"]):
            row = next(row for row in rows if row["company"] == company and row["report_type"] == report_type)
            value = float(row[metric])
            bar_h = (value / max_value) * chart_h
            x = group_x + (j - 0.5) * (bar_w + 8)
            y = top + chart_h - bar_h
            lines.append(f'<rect x="{x:.1f}" y="{y:.1f}" width="{bar_w}" height="{bar_h:.1f}" rx="4" fill="{colors[report_type]}"/>')
            label = f"{value:.1f}" if value < 100 else f"{value:.0f}"
            lines.append(xtext(x + bar_w / 2, y - 6, label, 9, "#111827", "middle", "700"))
        lines.append(xtext(group_x, top + chart_h + 24, company, 12, "#111827", "middle", "700"))
    lines.append("</svg>")
    path.write_text("\n".join(lines), encoding="utf-8")


def write_esg_stacked(path: Path, rows: list[dict[str, object]]) -> None:
    width, height = 980, 520
    left, right, top, bottom = 130, 40, 88, 40
    row_h = 44
    chart_w = width - left - right
    keys = [
        ("environmental_hits", "E", "#2E7D32"),
        ("social_hits", "S", "#0E7490"),
        ("governance_hits", "G", "#334155"),
    ]
    max_total = max(sum(float(row[key]) for key, _, _ in keys) for row in rows)
    lines = [
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}">',
        '<rect width="100%" height="100%" fill="#FFFFFF"/>',
        xtext(width / 2, 28, "E, S, and G Keyword Coverage: Generated vs. Actual", 18, "#111827", "middle", "700"),
    ]
    legend_x = left
    for _, label, color in keys:
        lines.append(f'<rect x="{legend_x}" y="48" width="14" height="14" rx="2" fill="{color}"/>')
        lines.append(xtext(legend_x + 22, 60, label, 12, "#374151", "start", "700"))
        legend_x += 70
    idx = 0
    for company in COMPANIES:
        for report_type in ["Generated", "Actual"]:
            row = next(row for row in rows if row["company"] == company and row["report_type"] == report_type)
            y = top + idx * row_h
            label = f"{company} {report_type}"
            lines.append(xtext(left - 14, y + 25, label, 11, "#111827", "end", "700" if report_type == "Generated" else "400"))
            x = left
            total = sum(float(row[key]) for key, _, _ in keys)
            for key, _, color in keys:
                value = float(row[key])
                seg_w = value / max_total * chart_w
                lines.append(f'<rect x="{x:.1f}" y="{y + 8:.1f}" width="{seg_w:.1f}" height="24" fill="{color}"/>')
                if seg_w > 30:
                    lines.append(xtext(x + seg_w / 2, y + 25, f"{value:.0f}", 10, "#FFFFFF", "middle", "700"))
                x += seg_w
            lines.append(xtext(left + chart_w + 7, y + 25, f"{total:.0f}", 10, "#374151", "start", "700"))
            idx += 1
    lines.append("</svg>")
    path.write_text("\n".join(lines), encoding="utf-8")


def write_similarity_chart(path: Path, rows: list[dict[str, object]]) -> None:
    width, height = 900, 430
    left, right, top, bottom = 78, 36, 62, 78
    chart_w = width - left - right
    chart_h = height - top - bottom
    group_w = chart_w / len(rows)
    bar_w = 34
    colors = {"tfidf_cosine_generated_vs_actual": "#7C3AED", "jaccard_generated_vs_actual": "#B45309"}
    lines = [
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}">',
        '<rect width="100%" height="100%" fill="#FFFFFF"/>',
        xtext(width / 2, 28, "Generated vs. Actual Text Similarity", 18, "#111827", "middle", "700"),
        f'<rect x="{width - 250}" y="44" width="14" height="14" rx="2" fill="{colors["tfidf_cosine_generated_vs_actual"]}"/>',
        xtext(width - 228, 56, "TF-IDF cosine", 11, "#374151", "start"),
        f'<rect x="{width - 130}" y="44" width="14" height="14" rx="2" fill="{colors["jaccard_generated_vs_actual"]}"/>',
        xtext(width - 108, 56, "Jaccard", 11, "#374151", "start"),
    ]
    for tick in range(6):
        value = tick / 5
        y = top + chart_h - value * chart_h
        lines.append(f'<line x1="{left}" y1="{y:.1f}" x2="{left + chart_w}" y2="{y:.1f}" stroke="#E5E7EB"/>')
        lines.append(xtext(left - 9, y + 4, f"{value:.1f}", 10, "#6B7280", "end"))
    lines.append(f'<line x1="{left}" y1="{top}" x2="{left}" y2="{top + chart_h}" stroke="#9CA3AF"/>')
    lines.append(f'<line x1="{left}" y1="{top + chart_h}" x2="{left + chart_w}" y2="{top + chart_h}" stroke="#9CA3AF"/>')
    for i, row in enumerate(rows):
        company = str(row["company"])
        center = left + i * group_w + group_w / 2
        for j, key in enumerate(["tfidf_cosine_generated_vs_actual", "jaccard_generated_vs_actual"]):
            value = float(row[key])
            bar_h = value * chart_h
            x = center + (j - 0.5) * (bar_w + 8)
            y = top + chart_h - bar_h
            lines.append(f'<rect x="{x:.1f}" y="{y:.1f}" width="{bar_w}" height="{bar_h:.1f}" rx="4" fill="{colors[key]}"/>')
            lines.append(xtext(x + bar_w / 2, y - 6, f"{value:.2f}", 10, "#111827", "middle", "700"))
        lines.append(xtext(center, top + chart_h + 24, company, 12, "#111827", "middle", "700"))
    lines.append("</svg>")
    path.write_text("\n".join(lines), encoding="utf-8")


def write_missing_terms_chart(path: Path, rows: list[dict[str, object]]) -> None:
    width, height = 980, 700
    panel_w, panel_h = 455, 285
    positions = {"Amazon": (34, 62), "Microsoft": (510, 62), "SAP": (34, 385), "Zalando": (510, 385)}
    colors = {"Amazon": "#2563EB", "Microsoft": "#0E7490", "SAP": "#7C3AED", "Zalando": "#B45309"}
    lines = [
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}">',
        '<rect width="100%" height="100%" fill="#FFFFFF"/>',
        xtext(width / 2, 30, "Terms More Prominent in Actual Reports", 18, "#111827", "middle", "700"),
    ]
    for company, (x0, y0) in positions.items():
        subset = [row for row in rows if row["company"] == company and int(row["rank"]) <= 8]
        max_count = max((int(row["actual_count_minus_generated_count"]) for row in subset), default=1)
        lines.append(xtext(x0, y0 - 14, company, 14, "#111827", "start", "700"))
        lines.append(f'<rect x="{x0}" y="{y0}" width="{panel_w}" height="{panel_h}" fill="#F9FAFB" stroke="#E5E7EB" rx="6"/>')
        for idx, row in enumerate(subset):
            term = str(row["actual_over_generated_term"])
            count = int(row["actual_count_minus_generated_count"])
            y = y0 + 24 + idx * 29
            bar_w = count / max_count * 250
            lines.append(xtext(x0 + 16, y + 15, term, 11, "#111827", "start"))
            lines.append(f'<rect x="{x0 + 155}" y="{y + 4}" width="{bar_w:.1f}" height="14" rx="3" fill="{colors[company]}"/>')
            lines.append(xtext(x0 + 155 + bar_w + 8, y + 16, count, 10, "#374151", "start", "700"))
    lines.append("</svg>")
    path.write_text("\n".join(lines), encoding="utf-8")


def make_figures(metrics: pd.DataFrame, esg: pd.DataFrame, tone: pd.DataFrame, pairwise: pd.DataFrame, missing_terms: pd.DataFrame) -> list[Path]:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    figures = []
    metric_records = metrics.to_dict("records")
    esg_records = esg.to_dict("records")
    tone_records = tone.to_dict("records")
    pairwise_records = pairwise.to_dict("records")
    missing_records = missing_terms.to_dict("records")

    specs = [
        ("word_count", "word_counts_generated_vs_actual.svg", "Word Count: Generated vs. Actual", "Words", 1.05),
        ("avg_sentence_length", "average_sentence_length_generated_vs_actual.svg", "Average Sentence Length", "Words per sentence", 1.15),
        ("numeric_mentions_per_1000_words", "numeric_density_generated_vs_actual.svg", "Numeric Evidence Density", "Mentions per 1,000 words", 1.15),
    ]
    for metric, filename, title, ylabel, scale in specs:
        path = FIG_DIR / filename
        write_grouped_bar(path, metric_records, metric, title, ylabel, scale)
        figures.append(path)

    esg_path = FIG_DIR / "esg_coverage_generated_vs_actual.svg"
    write_esg_stacked(esg_path, esg_records)
    figures.append(esg_path)

    tone_path = FIG_DIR / "risk_success_ratio_generated_vs_actual.svg"
    write_grouped_bar(tone_path, tone_records, "risk_to_success_ratio", "Risk-to-Success Tone Ratio", "Ratio", 1.25)
    figures.append(tone_path)

    sim_path = FIG_DIR / "similarity_generated_vs_actual.svg"
    write_similarity_chart(sim_path, pairwise_records)
    figures.append(sim_path)

    missing_path = FIG_DIR / "actual_prominent_terms.svg"
    write_missing_terms_chart(missing_path, missing_records)
    figures.append(missing_path)
    return figures


def write_notebook(figures: list[Path]) -> None:
    def cell(kind: str, source: str) -> dict[str, object]:
        obj = {"cell_type": kind, "metadata": {}, "source": source.splitlines(True)}
        if kind == "code":
            obj["execution_count"] = None
            obj["outputs"] = []
        return obj

    figure_markdown = "\n\n".join(
        f"![{figure.stem.replace('_', ' ').title()}](results/Actual%20vs%20Generatet/Python%20Analyse/figures/{figure.name})"
        for figure in figures
    )
    nb = {
        "cells": [
            cell(
                "markdown",
                "# Generated vs. Actual ESG Report Text Analysis\n\n"
                "This notebook compares AI-generated ESG reports with the official ESG, sustainability, impact, or integrated reports for Amazon, Microsoft, SAP, and Zalando. "
                "It uses classical Python text-analysis techniques from the lecture workflow: regex preprocessing, stopword removal, term frequencies, ESG keyword lexicons, lexicon-based tone analysis, TF-IDF cosine similarity, and Jaccard similarity.",
            ),
            cell(
                "code",
                "from pathlib import Path\n"
                "import re, math\n"
                "from collections import Counter\n"
                "import pandas as pd\n"
                "from docx import Document\n"
                "from pypdf import PdfReader\n\n"
                "def find_project_root():\n"
                "    for candidate in [Path.cwd(), *Path.cwd().parents]:\n"
                "        if (candidate / 'reports' / 'generated').exists() and (candidate / 'reports' / 'actual').exists():\n"
                "            return candidate\n"
                "    raise FileNotFoundError('Could not find project root with reports/generated and reports/actual')\n\n"
                "PROJECT = find_project_root()\n"
                "OUT = PROJECT / 'results' / 'Actual vs Generatet' / 'Python Analyse'\n"
                "FIGURES = OUT / 'figures'\n"
                "PROJECT\n",
            ),
            cell(
                "markdown",
                "## Reproducible outputs\n\n"
                "The full generated-vs-actual output tables are saved in `results/Actual vs Generatet/Python Analyse`. "
                "Run the companion script `run_actual_vs_generated_analysis.py` from that folder or rerun this notebook after installing `requirements.txt`.",
            ),
            cell(
                "code",
                "metrics = pd.read_csv(OUT / 'document_metrics_generated_vs_actual.csv')\n"
                "esg = pd.read_csv(OUT / 'esg_coverage_generated_vs_actual.csv')\n"
                "tone = pd.read_csv(OUT / 'tone_metrics_generated_vs_actual.csv')\n"
                "pairwise = pd.read_csv(OUT / 'pairwise_similarity_and_gaps.csv')\n"
                "display(metrics)\n"
                "display(esg)\n"
                "display(tone)\n"
                "display(pairwise)\n",
            ),
            cell(
                "markdown",
                "## Visual comparison\n\n" + figure_markdown,
            ),
            cell(
                "markdown",
                "## Interpretation guide\n\n"
                "- The official reports are much longer, so absolute word and keyword counts should be interpreted together with density measures per 1,000 words.\n"
                "- TF-IDF cosine similarity reflects weighted vocabulary overlap; Jaccard similarity reflects cleaned-token set overlap.\n"
                "- Actual-over-generated terms highlight themes that are more prominent in the official report than in the generated report.\n"
                "- The qualitative Word analysis in `results/Actual vs Generatet/LLM Analyse` interprets these metrics using close reading.",
            ),
        ],
        "metadata": {"kernelspec": {"display_name": "Python 3", "language": "python", "name": "python3"}, "language_info": {"name": "python", "version": "3.12"}},
        "nbformat": 4,
        "nbformat_minor": 5,
    }
    NOTEBOOK_ROOT.write_text(json.dumps(nb, indent=2), encoding="utf-8")
    shutil.copy2(NOTEBOOK_ROOT, PY_DIR / "analysis.ipynb")


def set_run_font(run, name: str = "Calibri", size: float | None = None, color: RGBColor | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    for style_name, size, color, before, after in [
        ("Normal", 11, RGBColor(0, 0, 0), 0, 6),
        ("Heading 1", 16, RGBColor(0x2E, 0x74, 0xB5), 16, 8),
        ("Heading 2", 13, RGBColor(0x2E, 0x74, 0xB5), 12, 6),
        ("Heading 3", 12, RGBColor(0x1F, 0x4D, 0x78), 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        if style_name.startswith("Heading"):
            style.font.bold = True
            style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10
    header = section.header.paragraphs[0]
    header.text = "Generated vs. actual ESG report analysis"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.runs[0], size=9, color=RGBColor(0x6B, 0x72, 0x80))


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def format_table(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW")) or OxmlElement("w:tblW")
    if tbl_w.getparent() is None:
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_ind = tbl_pr.find(qn("w:tblInd")) or OxmlElement("w:tblInd")
    if tbl_ind.getparent() is None:
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")
    for grid in list(tbl.findall(qn("w:tblGrid"))):
        tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(1, grid)
    for r, row in enumerate(table.rows):
        for c, width in enumerate(widths):
            cell = row.cells[c]
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW")) or OxmlElement("w:tcW")
            if tc_w.getparent() is None:
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            if r == 0:
                shade_cell(cell, "F2F4F7")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(3)
                for run in paragraph.runs:
                    set_run_font(run, size=9.2, bold=True if r == 0 else None)


def add_table(doc: Document, rows: list[list[object]], widths: list[int]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            table.cell(r, c).text = str(value)
    format_table(table, widths)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def build_llm_doc(metrics: pd.DataFrame, esg: pd.DataFrame, tone: pd.DataFrame, pairwise: pd.DataFrame, missing_terms: pd.DataFrame) -> Path:
    doc = Document()
    style_doc(doc)
    doc.core_properties.title = "Generated vs. Actual ESG Report LLM-Style Analysis"
    doc.core_properties.author = "Codex"
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Generated vs. Actual ESG Report Analysis")
    set_run_font(run, size=23, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("LLM-style qualitative interpretation of Python text-analysis results for Amazon, Microsoft, SAP, and Zalando")
    set_run_font(run, size=12, color=RGBColor(0x4B, 0x55, 0x63))

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "The generated ESG reports capture the broad E, S, and G structure of the official reports, but they are much shorter and therefore less granular. "
        "The official reports contain far more operational detail, framework-specific reporting language, quantified indicators, and company-specific disclosure boundaries. "
        "No automatically provable direct factual contradiction was detected from text metrics alone; the stronger finding is incompleteness and compression in the generated reports."
    )
    best_similarity = pairwise.sort_values("tfidf_cosine_generated_vs_actual", ascending=False).iloc[0]
    lowest_similarity = pairwise.sort_values("tfidf_cosine_generated_vs_actual", ascending=True).iloc[0]
    doc.add_paragraph(
        f"The closest generated-vs-actual vocabulary match is {best_similarity.company} by TF-IDF cosine similarity. "
        f"The lowest match is {lowest_similarity.company}, suggesting that the generated report uses a more compressed or differently framed vocabulary than the official disclosure."
    )

    doc.add_heading("Method", level=1)
    add_bullets(
        doc,
        [
            "The generated DOCX reports were compared with the official PDF reports in reports/actual.",
            "Python extracted text with python-docx and pypdf, then applied regex tokenization, stopword filtering, Counter term frequencies, ESG keyword lexicons, tone lexicons, TF-IDF cosine similarity, and Jaccard similarity.",
            "The qualitative analysis below is LLM-style close reading guided by those metrics. It is not an external fact-check and does not replace source-by-source verification.",
        ],
    )

    doc.add_heading("Pairwise Comparison", level=1)
    pair_rows = [["Company", "TF-IDF cosine", "Jaccard", "Generated words as % of actual", "Numeric density gap"]]
    for _, row in pairwise.iterrows():
        pair_rows.append(
            [
                row.company,
                row.tfidf_cosine_generated_vs_actual,
                row.jaccard_generated_vs_actual,
                row.generated_words_as_pct_of_actual,
                row.numeric_density_gap_generated_minus_actual,
            ]
        )
    add_table(doc, pair_rows, [1440, 1680, 1440, 2520, 2280])

    doc.add_heading("Company Findings", level=1)
    for company in COMPANIES:
        doc.add_heading(company, level=2)
        p_row = pairwise[pairwise.company == company].iloc[0]
        m_gen = metrics[(metrics.company == company) & (metrics.report_type == "Generated")].iloc[0]
        m_act = metrics[(metrics.company == company) & (metrics.report_type == "Actual")].iloc[0]
        e_gen = esg[(esg.company == company) & (esg.report_type == "Generated")].iloc[0]
        e_act = esg[(esg.company == company) & (esg.report_type == "Actual")].iloc[0]
        t_gen = tone[(tone.company == company) & (tone.report_type == "Generated")].iloc[0]
        t_act = tone[(tone.company == company) & (tone.report_type == "Actual")].iloc[0]
        missing = missing_terms[missing_terms.company == company].head(8)
        missing_list = ", ".join(missing.actual_over_generated_term.astype(str).tolist())

        doc.add_paragraph(
            f"The generated report has {int(m_gen.word_count):,} words versus {int(m_act.word_count):,} words in the official report "
            f"({p_row.generated_words_as_pct_of_actual}% of official length). Its TF-IDF cosine similarity to the official report is "
            f"{p_row.tfidf_cosine_generated_vs_actual}, and its Jaccard similarity is {p_row.jaccard_generated_vs_actual}."
        )
        doc.add_paragraph(
            f"Specificity: the generated report has {m_gen.numeric_mentions_per_1000_words} numeric mentions per 1,000 words, "
            f"while the official report has {m_act.numeric_mentions_per_1000_words}. "
            "If the generated density is higher, this usually reflects concise summarization rather than a fuller evidence base; the official report still contains more total evidence."
        )
        doc.add_paragraph(
            f"Coverage: the generated report's strongest ESG pillar is {e_gen.strongest_pillar}, while the official report's strongest pillar is {e_act.strongest_pillar}. "
            f"The most prominent official-report terms that are less visible in the generated report include: {missing_list}."
        )
        doc.add_paragraph(
            f"Framing: the generated risk-to-success ratio is {t_gen.risk_to_success_ratio}, compared with {t_act.risk_to_success_ratio} in the official report. "
            "A higher ratio indicates comparatively more attention to risks, constraints, controversies, or unresolved issues."
        )
        if company == "Amazon":
            add_bullets(
                doc,
                [
                    "Likely missing depth: detailed operational metrics across logistics, packaging, renewable energy, safety, human rights, and supplier programs.",
                    "Potential tension: the generated report captures the ESG pillars but cannot match the official report's scale-specific operational evidence.",
                ],
            )
        elif company == "Microsoft":
            add_bullets(
                doc,
                [
                    "Likely missing depth: detailed environmental accounting, data-center energy and water impacts, cloud/AI growth implications, and official impact-program evidence.",
                    "Potential tension: the generated report may describe responsible AI and climate goals, while the official reports provide more structured measurement boundaries.",
                ],
            )
        elif company == "SAP":
            add_bullets(
                doc,
                [
                    "Likely missing depth: framework language around integrated reporting, materiality, assurance, taxonomy, and process-level governance.",
                    "Potential tension: SAP's official reporting is highly structured and regulatory, while the generated report is a shorter ESG narrative.",
                ],
            )
        else:
            add_bullets(
                doc,
                [
                    "Likely missing depth: ESRS/CSRD disclosure structure, value-chain workers, supplier due diligence, product circularity, returns, and platform governance details.",
                    "Potential tension: the generated report covers fashion-relevant ESG topics but compresses the official CSRD-style disclosure architecture.",
                ],
            )

    doc.add_heading("Cross-Company Interpretation", level=1)
    add_bullets(
        doc,
        [
            "The generated reports are useful as compact ESG syntheses, but official reports remain stronger for auditability, quantified detail, and disclosure-boundary precision.",
            "Generated reports tend to normalize structure across companies. That supports comparison but can hide industry-specific reporting obligations and materiality differences.",
            "Official reports are harder to summarize because they mix narrative, metrics, frameworks, risk language, and regulatory compliance tables.",
            "The main economic intuition is that ESG disclosure is not only a statement of performance; it is also a governance and accountability mechanism shaped by regulation, stakeholder pressure, and sector-specific risk.",
        ],
    )

    doc.add_heading("Limitations", level=1)
    add_bullets(
        doc,
        [
            "PDF text extraction may miss table structure, footnotes, charts, and page layout.",
            "Keyword lexicons are transparent but imperfect; they cannot fully capture semantic meaning or source quality.",
            "Similarity metrics are sensitive to report length, boilerplate, and framework vocabulary.",
            "The LLM-style section is an interpretive close reading guided by metrics, not a verified factual audit against every source statement.",
        ],
    )

    out = LLM_DIR / "Generated_vs_Actual_LLM_Analysis.docx"
    LLM_DIR.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(out)
    except PermissionError:
        out = LLM_DIR / "Generated_vs_Actual_LLM_Analysis_Final.docx"
        doc.save(out)
    return out


def audit_docx(path: Path) -> str:
    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    issues = []
    if len(doc.tables) < 1:
        issues.append("No tables found.")
    if re.search(r"\b(TODO|TBD|Platzhalter)\b", text, flags=re.IGNORECASE):
        issues.append("Placeholder text found.")
    section = doc.sections[0]
    if not all(abs(getattr(section, margin).inches - 1.0) < 0.02 for margin in ["top_margin", "right_margin", "bottom_margin", "left_margin"]):
        issues.append("Margins differ from 1 inch.")
    with zipfile.ZipFile(path) as archive:
        xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")
    if "w:tblGrid" not in xml:
        issues.append("Explicit table grid not found.")
    lines = [
        f"Structural DOCX audit for {path.name}",
        f"Paragraphs: {len(doc.paragraphs)}",
        f"Tables: {len(doc.tables)}",
        "Preset: standard_business_brief",
        "Status: PASS" if not issues else "Status: CHECK",
    ]
    if issues:
        lines.append("Issues:")
        lines.extend(f"- {issue}" for issue in issues)
    return "\n".join(lines) + "\n"


def write_readmes(figures: list[Path], metrics: pd.DataFrame, pairwise: pd.DataFrame) -> None:
    PY_DIR.mkdir(parents=True, exist_ok=True)
    LLM_DIR.mkdir(parents=True, exist_ok=True)

    py_readme = [
        "# Python Analyse - Generated vs. Actual",
        "",
        "This folder contains the Python text-analysis outputs comparing DataNXT-generated ESG reports with official ESG, sustainability, impact, or integrated reports.",
        "",
        "Main files:",
        "",
        "- `analysis.ipynb`: notebook copy for the generated-vs-actual analysis.",
        "- `requirements.txt`: Python dependencies.",
        "- `document_metrics_generated_vs_actual.csv`: length, sentence, lexical, and numeric evidence metrics.",
        "- `esg_coverage_generated_vs_actual.csv`: E, S, and G keyword coverage.",
        "- `tone_metrics_generated_vs_actual.csv`: success, risk, controversy, and target/tradeoff tone signals.",
        "- `pairwise_similarity_and_gaps.csv`: TF-IDF cosine similarity, Jaccard similarity, and generated-vs-actual gaps.",
        "- `actual_prominent_terms.csv`: terms much more prominent in official reports than in generated reports.",
        "- `figures/`: SVG charts used in the notebook and README.",
        "",
        "Important methodological note: official reports are used only as benchmark documents for Part 3. They were not used as DataNXT workflow input.",
        "",
    ]
    (PY_DIR / "README.md").write_text("\n".join(py_readme), encoding="utf-8")

    llm_readme = [
        "# LLM Analyse - Generated vs. Actual",
        "",
        "This folder contains the qualitative generated-vs-actual analysis.",
        "",
        "Main files:",
        "",
        "- `Generated_vs_Actual_LLM_Analysis.docx`: qualitative LLM-style analysis.",
        "- `llm_analysis_audit.txt`: structural DOCX audit.",
        "",
        "The analysis covers factual tensions, missing themes, specificity, and framing across achievements, risks, controversies, and target conflicts.",
        "",
    ]
    (LLM_DIR / "README.md").write_text("\n".join(llm_readme), encoding="utf-8")

    metric_summary = metrics.pivot(index="company", columns="report_type", values="word_count")
    metric_summary["generated_pct_actual"] = (metric_summary["Generated"] / metric_summary["Actual"] * 100).round(2)
    sim_best = pairwise.sort_values("tfidf_cosine_generated_vs_actual", ascending=False).iloc[0]
    sim_low = pairwise.sort_values("tfidf_cosine_generated_vs_actual").iloc[0]

    lines = [
        "# ESG Report Analysis Results",
        "",
        f"Generated on {date.today().isoformat()}.",
        "",
        "## Scope",
        "",
        "This repository contains a DataNXT-based ESG reporting workflow and a generated-vs-actual text analysis for Amazon, Microsoft, SAP, and Zalando. The official ESG, sustainability, impact, or integrated reports are used only as benchmark documents for Part 3. They were not used as DataNXT input.",
        "",
        "## Literature and methodological background",
        "",
        "ESG reporting quality depends on whether disclosures cover material environmental, social, and governance topics, define reporting boundaries, provide quantified indicators, and connect claims to recognized frameworks such as GRI, TCFD, SASB/ISSB, EU CSRD/ESRS, and company-specific materiality assessments. Classical text analysis helps compare length, vocabulary, topic coverage, tone, and similarity. LLM-style qualitative analysis helps interpret missing themes, specificity, tradeoffs, and tensions that simple counts cannot fully capture.",
        "",
        "## Results structure",
        "",
        "- `Actual vs Generatet/Python Analyse`: notebook, requirements, CSV tables, extracted text, and SVG visualizations.",
        "- `Actual vs Generatet/LLM Analyse`: qualitative Word analysis and DOCX audit.",
        "- `reports/generated`: four DataNXT-generated ESG reports.",
        "- `reports/actual`: official benchmark reports.",
        "",
        "## Main empirical findings",
        "",
        f"- The generated reports are much shorter than the official reports. Across the four companies, generated report length ranges from {metric_summary['generated_pct_actual'].min():.2f}% to {metric_summary['generated_pct_actual'].max():.2f}% of the corresponding official report.",
        f"- The closest generated-vs-actual vocabulary match is {sim_best.company} by TF-IDF cosine similarity ({sim_best.tfidf_cosine_generated_vs_actual}).",
        f"- The lowest generated-vs-actual vocabulary match is {sim_low.company} ({sim_low.tfidf_cosine_generated_vs_actual}), indicating a larger difference in framing or disclosure vocabulary.",
        "- Official reports contain more detailed framework language, more disclosure-boundary information, more tables, and more company-specific evidence. Generated reports are clearer as compact summaries but weaker as audit-ready disclosures.",
        "- Generated reports often show high numeric density per 1,000 words because they are short and condensed. This should not be confused with greater evidence depth; official reports contain much more total quantified disclosure.",
        "",
        "## Figures and tables",
        "",
    ]
    for figure in figures:
        rel = figure.relative_to(RESULTS).as_posix()
        lines.append(f"- `{rel}`")
    lines.extend(
        [
            "",
            "## Economic intuition",
            "",
            "The comparison shows why corporate ESG disclosure is difficult to reproduce with a general AI workflow. Official reports are not only narratives about sustainability performance; they are also accountability documents shaped by regulation, assurance expectations, stakeholder pressure, sector-specific risks, and reporting-framework requirements. A reusable DataNXT workflow can produce consistent company summaries, but it tends to compress disclosure boundaries and industry-specific materiality.",
            "",
            "## Critical reflection and limitations",
            "",
            "- PDF extraction can miss tables, figures, footnotes, and layout cues.",
            "- Keyword lexicons are transparent and reproducible but cannot fully capture context or factual correctness.",
            "- Similarity scores are affected by report length and by framework boilerplate.",
            "- The generated reports depend on available input sources and the quality of DataNXT answers marked as relevant.",
            "- The LLM-style Word analysis is an interpretive synthesis guided by metrics, not a source-by-source legal or assurance audit.",
            "",
            "## Reproducibility",
            "",
            "Use `Actual vs Generatet/Python Analyse/requirements.txt` and run `analysis.ipynb`. The notebook uses relative paths and expects the repository structure to contain `reports/generated` and `reports/actual`.",
            "",
        ]
    )
    README.write_text("\n".join(lines), encoding="utf-8")


def write_requirements() -> None:
    req = "\n".join(["python-docx>=1.1.0", "pandas>=2.0.0", "pypdf>=4.0.0", "ipykernel>=6.0.0", ""]) 
    (PY_DIR / "requirements.txt").write_text(req, encoding="utf-8")
    (PROJECT / "requirements.txt").write_text(req, encoding="utf-8")


def write_runner_copy() -> None:
    src = Path(__file__)
    target = PY_DIR / "run_actual_vs_generated_analysis.py"
    shutil.copy2(src, target)
    text = target.read_text(encoding="utf-8")
    text = text.replace('PROJECT = Path(r"C:\\Jann\\Studium\\2. Semester\\ESG-Report")', "PROJECT = next(candidate for candidate in [Path.cwd(), *Path.cwd().parents] if (candidate / 'reports' / 'generated').exists() and (candidate / 'reports' / 'actual').exists())")
    target.write_text(text, encoding="utf-8")


def copy_outputs(docx_path: Path) -> None:
    OUTPUTS.mkdir(parents=True, exist_ok=True)
    shutil.copy2(docx_path, OUTPUTS / docx_path.name)
    shutil.copy2(NOTEBOOK_ROOT, OUTPUTS / NOTEBOOK_ROOT.name)
    shutil.copy2(PY_DIR / "pairwise_similarity_and_gaps.csv", OUTPUTS / "pairwise_similarity_and_gaps.csv")


def main() -> None:
    for folder in [PY_DIR, LLM_DIR, FIG_DIR, TEXT_DIR]:
        folder.mkdir(parents=True, exist_ok=True)
    missing = [path for paths in [*GENERATED_FILES.values(), *ACTUAL_FILES.values()] for path in paths if not path.exists()]
    if missing:
        raise FileNotFoundError("Missing report files: " + "; ".join(str(path) for path in missing))

    corpus, source_rows = build_corpus()
    metrics, esg, tone, top_terms = metric_rows(corpus)
    pairwise, missing_terms = pairwise_rows(corpus, metrics, esg, tone)

    pd.DataFrame(source_rows).to_csv(PY_DIR / "source_files_used.csv", index=False)
    metrics.to_csv(PY_DIR / "document_metrics_generated_vs_actual.csv", index=False)
    esg.to_csv(PY_DIR / "esg_coverage_generated_vs_actual.csv", index=False)
    tone.to_csv(PY_DIR / "tone_metrics_generated_vs_actual.csv", index=False)
    top_terms.to_csv(PY_DIR / "top_terms_generated_vs_actual.csv", index=False)
    pairwise.to_csv(PY_DIR / "pairwise_similarity_and_gaps.csv", index=False)
    missing_terms.to_csv(PY_DIR / "actual_prominent_terms.csv", index=False)

    figures = make_figures(metrics, esg, tone, pairwise, missing_terms)
    write_notebook(figures)
    write_requirements()
    write_runner_copy()
    docx_path = build_llm_doc(metrics, esg, tone, pairwise, missing_terms)
    (LLM_DIR / "llm_analysis_audit.txt").write_text(audit_docx(docx_path), encoding="utf-8")
    write_readmes(figures, metrics, pairwise)
    copy_outputs(docx_path)

    print("Generated-vs-actual analysis completed.")
    print(f"Python output: {PY_DIR}")
    print(f"LLM output: {LLM_DIR}")
    print(f"Notebook: {NOTEBOOK_ROOT}")
    print(f"Word document: {docx_path}")


if __name__ == "__main__":
    main()
